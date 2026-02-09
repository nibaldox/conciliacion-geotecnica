
import io
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import numpy as np

def create_section_plot(params_design, params_topo, distances_d, elevations_d, distances_t, elevations_t):
    """
    Generate a matplotlib figure for a section comparison.
    Overlay detected benches (Crests/Toes) to show 'Reconciled Profile'.
    Returns: BytesIO object containing the image.
    """
    fig, ax = plt.subplots(figsize=(12, 7))
    
    # 1. Plot Design (Baseline)
    if len(distances_d) > 0:
        ax.plot(distances_d, elevations_d, 'c-', label='Diseño', linewidth=1.5, alpha=0.7)
        
    # 2. Plot As-Built (Real)
    if len(distances_t) > 0:
        ax.plot(distances_t, elevations_t, 'r-', label='Perfil Real (Scan)', linewidth=2, alpha=0.6)
        
    # 3. Plot Reconciled Segments (Crest -> Toe -> Crest...)
    # We reconstruct the "interpreted" profile from the extracted parameters
    if params_topo.benches:
        rec_dist = []
        rec_elev = []
        
        # Collect points: Toe_i -> Crest_i -> Toe_i+1 ...
        # This is an approximation of the "Green Line" in the user's image
        for b in params_topo.benches:
            # Toe
            rec_dist.append(b.toe_point[0])
            rec_elev.append(b.toe_point[1])
            # Crest
            rec_dist.append(b.crest_point[0])
            rec_elev.append(b.crest_point[1])
            
            # Draw markers
            ax.plot(b.toe_point[0], b.toe_point[1], 'go', markersize=4) # Toe marker
            ax.plot(b.crest_point[0], b.crest_point[1], 'go', markersize=4) # Crest marker
            
            # Annotate?
            # ax.text(b.crest_point[0], b.crest_point[1], f"C{b.bench_number}", fontsize=8)

        if rec_dist:
            ax.plot(rec_dist, rec_elev, 'g-', label='Perfil Conciliado', linewidth=2.5)

    # Highlight Design Benches too for reference
    if params_design.benches:
        for b in params_design.benches:
            ax.plot(b.crest_point[0], b.crest_point[1], 'c.', markersize=6)

    ax.set_title(f"Sección: {params_design.section_name} - {params_design.sector}")
    ax.set_xlabel("Distancia (m)")
    ax.set_ylabel("Elevación (m)")
    ax.grid(True, linestyle='--', alpha=0.5)
    ax.legend(loc='upper right')
    ax.set_aspect('equal', adjustable='box')
    
    # Save to buffer
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf

def generate_word_report(comparisons, all_data, output_path, project_info=None):
    """
    Generate a full Word report with summary tables and section detailed plots.
    
    all_data: List of dicts containing keys:
        - section_name
        - params_design
        - params_topo
        - profile_d: (dist, elev) arrays
        - profile_t: (dist, elev) arrays
    """
    if project_info is None:
        project_info = {}
        
    doc = Document()
    
    # Title
    title = doc.add_heading(f"Informe de Conciliación Geotécnica", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Meta info
    p = doc.add_paragraph()
    p.add_run(f"Proyecto: {project_info.get('project', 'N/A')}\n").bold = True
    p.add_run(f"Elaborado por: {project_info.get('author', 'N/A')}\n")
    p.add_run(f"Fecha: {datetime.now().strftime('%d/%m/%Y')}\n")
    
    # Summary
    doc.add_heading("1. Resumen Ejecutivo", level=1)
    
    if comparisons:
        n_total = len(comparisons) * 3
        n_ok = sum(1 for c in comparisons for k in ['height_status','angle_status','berm_status'] if c[k] == "CUMPLE")
        pct = n_ok / n_total * 100 if n_total > 0 else 0
        
        doc.add_paragraph(f"Se evaluaron {len(comparisons)} bancos en total.")
        doc.add_paragraph(f"Cumplimiento Global: {pct:.1f}%")
        
        # Table of Compliance
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Parámetro'
        hdr_cells[1].text = 'CUMPLE'
        hdr_cells[2].text = 'FUERA TOL.'
        hdr_cells[3].text = 'NO CUMPLE'
        
        for key, label in [('height_status', 'Altura'), ('angle_status', 'Ángulo Cara'), ('berm_status', 'Berma')]:
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = str(sum(1 for c in comparisons if c[key] == "CUMPLE"))
            row_cells[2].text = str(sum(1 for c in comparisons if c[key] == "FUERA DE TOLERANCIA"))
            row_cells[3].text = str(sum(1 for c in comparisons if c[key] == "NO CUMPLE"))
    else:
        doc.add_paragraph("No se encontraron resultados para reportar.")

    # Detailed Sections
    doc.add_heading("2. Detalle por Sección", level=1)
    
    for item in all_data:
        sec_name = item['section_name']
        doc.add_heading(f"Sección {sec_name}", level=2)
        
        # Add Plot
        pd = item['params_design']
        pt = item['params_topo']
        prof_d = item['profile_d']
        prof_t = item['profile_t']
        
        img_stream = create_section_plot(pd, pt, prof_d[0], prof_d[1], prof_t[0], prof_t[1])
        doc.add_picture(img_stream, width=Inches(6))
        img_stream.close()
        
        # Add bench table for this section
        sec_comps = [c for c in comparisons if c['section'] == sec_name]
        if sec_comps:
            table = doc.add_table(rows=1, cols=7)
            table.style = 'Table Grid'
            # Header
            headers = ['Banco', 'H. Dise (m)', 'H. Real', 'Ang. Dise (°)', 'Ang. Real', 'Berma Real (m)', 'Estado']
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
                
            for c in sec_comps:
                row_cells = table.add_row().cells
                row_cells[0].text = str(c['bench_num'])
                row_cells[1].text = str(c['height_design'])
                row_cells[2].text = str(c['height_real'])
                row_cells[3].text = str(c['angle_design'])
                row_cells[4].text = str(c['angle_real'])
                row_cells[5].text = str(c['berm_real'])
                
                # Check overall status for this bench (simple logic: all must be compliant)
                statuses = [c['height_status'], c['angle_status'], c['berm_status']]
                if "NO CUMPLE" in statuses:
                    final_status = "NO CUMPLE"
                elif "FUERA DE TOLERANCIA" in statuses:
                    final_status = "ALERTA"
                else:
                    final_status = "OK"
                row_cells[6].text = final_status

        doc.add_page_break()

    doc.save(output_path)


def generate_section_images_zip(all_data):
    """
    Generate a ZIP file containing PNG images of all section plots.
    Returns: BytesIO object containing the ZIP file.
    """
    import zipfile
    
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
        for item in all_data:
            sec_name = item['section_name']
            
            # Extract data
            pd = item['params_design']
            pt = item['params_topo']
            prof_d = item['profile_d']
            prof_t = item['profile_t']
            
            # Generate plot
            img_buf = create_section_plot(pd, pt, prof_d[0], prof_d[1], prof_t[0], prof_t[1])
            
            # Add to ZIP
            filename = f"{sec_name}.png"
            zip_file.writestr(filename, img_buf.getvalue())
            img_buf.close()
            
    zip_buffer.seek(0)
    return zip_buffer
